#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор звіту про оцінку комерційної нерухомості.
Методологія: .agents/skills/cre-valuation/ [B1] [B2] [B3]

Режим Full Data Audit: перевіряє 14 обов'язкових параметрів, інтерактивно
запитує відсутні та оновлює YAML-frontmatter у MD-картці об'єкта.

Підходи:
  [B1] Порівняльний — матриця корегувань для 5+ аналогів
  [B2] Дохідний    — PGI → EGI → NOI → Cap Rate (Excel-формули)
  [B3] Узгодження  — зважена середня (за замовч. 50/50)

Запуск:
    uv run agents/valuation_report.py --name "Фастов завод" \
        --area 5000 --type Warehouse --rent-rate 5.5

    uv run agents/valuation_report.py --help
"""

from __future__ import annotations

import argparse
import re
import sys
from collections import OrderedDict
from datetime import date
from pathlib import Path
from typing import Any

import yaml

# ─── Залежності ───────────────────────────────────────────────────────────────
try:
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    sys.exit("Встановіть: pip3 install openpyxl")

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    sys.exit("Встановіть: pip3 install python-docx")

# ─── Шляхи ────────────────────────────────────────────────────────────────────
BASE_DIR      = Path(__file__).parent.parent
OBJECTS_DIR   = BASE_DIR / "Объекты"
CLIPPINGS_DIR = BASE_DIR / "Clippings"
OUTPUT_DIR    = BASE_DIR / "output" / "reports"

# ─── Full Data Audit — обов'язковий список параметрів ─────────────────────────
# Поля: YAML-ключ → {label, type}
REQUIRED_AUDIT_FIELDS: OrderedDict = OrderedDict([
    ("District",            {"label": "Район міста (напр. Шевченківський)",             "type": "str"}),
    ("Location_Zone",       {"label": "Зона локації (Center/Middle/Periphery/Suburbs)",  "type": "str"}),
    ("Distance_to_Metro",   {"label": "Відстань до метро, хв пішки (число)",            "type": "float"}),
    ("Building_Class",      {"label": "Клас будівлі (A / B+ / B / C)",                  "type": "str"}),
    ("Floor_Type",          {"label": "Тип поверху (1й поверх / мансарда / підвал)",    "type": "str"}),
    ("Area",                {"label": "Загальна площа GBA, м²",                         "type": "float"}),
    ("Ceiling_Height",      {"label": "Висота стелі, м",                                "type": "float"}),
    ("Condition_Type",      {"label": "Стан об'єкта (з ремонтом / під оздоблення / без ремонту)", "type": "str"}),
    ("Renovation_Style",    {"label": "Стиль ремонту (представницький / офісний / без оздоблення)", "type": "str"}),
    ("Generator",           {"label": "Генератор є? (так / ні)",                        "type": "bool"}),
    ("Shelter",             {"label": "Укриття є? (так / ні)",                          "type": "bool"}),
    ("Parking_Underground", {"label": "Підземний паркінг — кількість місць (0 якщо немає)", "type": "int"}),
    ("Parking_Open",        {"label": "Відкритий/наземний паркінг — кількість місць",   "type": "int"}),
    ("OPEX_Owner",          {"label": "OPEX — на кому (орендар / власник / змішано)",   "type": "str"}),
])

# Mapping YAML-ключ → ключ у subject dict
_AUDIT_FIELD_MAP: dict[str, str] = {
    "District":            "district",
    "Location_Zone":       "zone",
    "Distance_to_Metro":   "metro_min",
    "Building_Class":      "building_class",
    "Floor_Type":          "floor_type",
    "Area":                "area",
    "Ceiling_Height":      "ceiling_height",
    "Condition_Type":      "condition",
    "Renovation_Style":    "renovation_style",
    "Generator":           "generator",
    "Shelter":             "shelter",
    "Parking_Underground": "parking_underground",
    "Parking_Open":        "parking_open",
    "OPEX_Owner":          "opex_owner",
}

# ─── Ринкові константи ─────────────────────────────────────────────────────────
DISCOUNT_RATES: dict[str, float] = {
    "Warehouse":   0.07,
    "Office":      0.08,
    "Retail":      0.06,
    "Land":        0.05,
    "Residential": 0.05,
}

CAP_RATES: dict[str, float] = {
    "Warehouse":   0.12,
    "Office":      0.10,
    "Retail":      0.11,
    "Land":        0.08,
    "Residential": 0.08,
}

OPEX_PCT: dict[str, float] = {
    "Warehouse":   0.15,
    "Office":      0.20,
    "Retail":      0.18,
    "Land":        0.05,
    "Residential": 0.25,
}

ZONE_RANK: dict[str, int] = {
    "Center": 4, "Middle": 3, "Periphery": 2, "Suburbs": 1, "Unknown": 2,
}

# ─── Шаблонні аналоги ──────────────────────────────────────────────────────────
DEMO_ANALOGS: dict[str, list[dict]] = {
    "Warehouse": [
        {"name": "Склад, Дарниця",              "area": 5000,  "price_psm": 270,
         "zone": "Periphery", "condition": "з ремонтом",  "ceiling": 4.0, "railway": True},
        {"name": "Вир.-склад. пр., Радистів",   "area": 1500,  "price_psm": 320,
         "zone": "Periphery", "condition": "з ремонтом",  "ceiling": 6.0, "railway": False},
        {"name": "Склад. комплекс, Мила",        "area": 10323, "price_psm": 242,
         "zone": "Suburbs",   "condition": "з ремонтом",  "ceiling": 7.0, "railway": False},
        {"name": "Склад, Видубичі",              "area": 3000,  "price_psm": 933,
         "zone": "Center",    "condition": "з ремонтом",  "ceiling": 5.0, "railway": False},
        {"name": "Логіст. центр, Бровари",       "area": 8000,  "price_psm": 150,
         "zone": "Suburbs",   "condition": "після будівельників", "ceiling": 10.0, "railway": True},
    ],
    "Office": [
        {"name": "Офіс, Хмельницького, 750м²",  "area": 750,  "price_psm": 5000,
         "zone": "Center",  "condition": "з ремонтом", "class": "A",  "metro_min": 1},
        {"name": "БЦ, Голосіїв, 1200м²",        "area": 1200, "price_psm": 3000,
         "zone": "Middle",  "condition": "з ремонтом", "class": "B+", "metro_min": 7},
        {"name": "Офіс, Поділ, 500м²",          "area": 500,  "price_psm": 6000,
         "zone": "Center",  "condition": "з ремонтом", "class": "B+", "metro_min": 5},
        {"name": "БЦ, Лук'янівка, 2000м²",      "area": 2000, "price_psm": 2500,
         "zone": "Middle",  "condition": "з ремонтом", "class": "B",  "metro_min": 10},
        {"name": "Офіс, Оболонь, 900м²",        "area": 900,  "price_psm": 2200,
         "zone": "Middle",  "condition": "під оздоблення", "class": "B", "metro_min": 8},
    ],
    "Retail": [
        {"name": "Магазин, Хрещатик, 300м²",    "area": 300,  "price_psm": 13000,
         "zone": "Center",    "condition": "з ремонтом"},
        {"name": "Рітейл, Оболонь, 450м²",      "area": 450,  "price_psm": 5000,
         "zone": "Middle",    "condition": "з ремонтом"},
        {"name": "Магазин, Троєщина, 600м²",    "area": 600,  "price_psm": 3000,
         "zone": "Periphery", "condition": "з ремонтом"},
        {"name": "Стріт-рітейл, Поділ, 200м²",  "area": 200,  "price_psm": 13000,
         "zone": "Center",    "condition": "з ремонтом"},
        {"name": "Рітейл, Бровари, 800м²",      "area": 800,  "price_psm": 2000,
         "zone": "Suburbs",   "condition": "під оздоблення"},
    ],
}


# ─────────────────────────────────────────────────────────────────────────────
# Full Data Audit — перевірка і доповнення параметрів
# ─────────────────────────────────────────────────────────────────────────────

FRONTMATTER_RE = re.compile(r"^---\n(.*?)\n---", re.DOTALL)


def _cast_value(raw: str, field_type: str) -> Any:
    """Приведення відповіді користувача до потрібного типу."""
    raw = raw.strip()
    if field_type == "float":
        return float(raw.replace(",", "."))
    if field_type == "int":
        return int(float(raw.replace(",", ".")))
    if field_type == "bool":
        return raw.lower() in ("так", "yes", "y", "true", "1", "+")
    return raw


def find_object_md(obj_name: str) -> Path | None:
    """Шукає головну MD-картку в Объекты/{obj_name}/wiki/objects/."""
    obj_dir = OBJECTS_DIR / obj_name
    if not obj_dir.exists():
        return None
    objects_wiki = obj_dir / "wiki" / "objects"
    if not objects_wiki.exists():
        return None
    mds = sorted(objects_wiki.glob("*.md"),
                 key=lambda p: p.stat().st_mtime, reverse=True)
    return mds[0] if mds else None


def load_yaml_from_md(md_path: Path) -> dict:
    """Читає YAML-frontmatter з MD-файлу."""
    try:
        text = md_path.read_text(encoding="utf-8")
        m = FRONTMATTER_RE.match(text)
        if m:
            return yaml.safe_load(m.group(1)) or {}
    except Exception:
        pass
    return {}


def update_yaml_in_md(md_path: Path, updates: dict) -> None:
    """Оновлює YAML-frontmatter у MD-файлі, зберігаючи решту тексту."""
    text  = md_path.read_text(encoding="utf-8")
    m     = FRONTMATTER_RE.match(text)
    if m:
        old_data = yaml.safe_load(m.group(1)) or {}
        old_data.update(updates)
        new_fm = yaml.dump(old_data, allow_unicode=True,
                           default_flow_style=False, sort_keys=False).rstrip()
        new_text = f"---\n{new_fm}\n---" + text[m.end():]
    else:
        fm_str = yaml.dump(updates, allow_unicode=True,
                           default_flow_style=False, sort_keys=False).rstrip()
        new_text = f"---\n{fm_str}\n---\n\n" + text
    md_path.write_text(new_text, encoding="utf-8")


def full_data_audit(subject: dict, obj_name: str | None = None,
                    md_path: Path | None = None) -> dict:
    """
    Full Data Audit [FDA]:
      1. Зчитує наявні дані з YAML-frontmatter (якщо є MD-файл)
      2. Визначає відсутні з 14 обов'язкових полів
      3. Інтерактивно запитує кожне відсутнє поле
      4. Оновлює YAML-frontmatter у MD-файлі
      5. Повертає оновлений subject dict
    """
    # 1. Зчитуємо з YAML, якщо є
    if md_path and md_path.exists():
        yaml_data = load_yaml_from_md(md_path)
        for yaml_key, subj_key in _AUDIT_FIELD_MAP.items():
            if subject.get(subj_key) is None:
                val = yaml_data.get(yaml_key)
                if val is not None:
                    subject[subj_key] = val

    # 2. Знаходимо відсутні поля
    missing = [
        (yaml_key, _AUDIT_FIELD_MAP[yaml_key], meta)
        for yaml_key, meta in REQUIRED_AUDIT_FIELDS.items()
        if subject.get(_AUDIT_FIELD_MAP[yaml_key]) is None
    ]

    if not missing:
        print("  ✓ Full Data Audit: всі 14 обов'язкових параметрів присутні.")
        return subject

    name_display = obj_name or subject.get("name", "об'єкт")
    print(f"\n{'='*62}")
    print(f"  ВНИМАНИЕ: Недостаточно данных для объекта «{name_display}»")
    print(f"  Відсутніх параметрів: {len(missing)} / {len(REQUIRED_AUDIT_FIELDS)}")
    print(f"{'='*62}")

    # 3. Інтерактивне заповнення
    collected: dict = {}
    for yaml_key, subj_key, meta in missing:
        label = meta["label"]
        ftype = meta["type"]
        while True:
            try:
                raw = input(f"\n  [{yaml_key}] {label}: ").strip()
                if not raw:
                    print("  ⚠ Пропущено (залишається порожнім).")
                    break
                val = _cast_value(raw, ftype)
                subject[subj_key] = val
                collected[yaml_key] = val
                break
            except (ValueError, KeyboardInterrupt):
                print("  ✗ Невірний формат. Спробуйте ще раз.")

    # 4. Оновлюємо MD-файл
    if md_path and collected:
        try:
            update_yaml_in_md(md_path, collected)
            print(f"\n  ✓ YAML оновлено: {md_path.name}")
        except Exception as e:
            print(f"\n  ⚠ Не вдалося оновити {md_path.name}: {e}")

    print()
    return subject


def ask_dynamic_inputs(obj_type: str) -> tuple[float, float]:
    """
    Перед генерацією запитує:
      1. Знижку на торг (%)
      2. Ставку капіталізації Cap Rate (%)

    Returns:
        (discount_rate, cap_rate) — дробові числа 0.0–1.0
    """
    default_disc = DISCOUNT_RATES.get(obj_type, 0.07)
    default_cap  = CAP_RATES.get(obj_type, 0.11)

    print(f"\n{'─'*52}")
    print("  ДИНАМІЧНІ ВХІДНІ ПАРАМЕТРИ")
    print(f"{'─'*52}")

    # Знижка на торг
    while True:
        raw = input(
            f"  Розмір знижки на торг (%) "
            f"[орієнтир {default_disc*100:.0f}%, Enter = прийняти]: "
        ).strip()
        if not raw:
            discount_rate = default_disc
            break
        try:
            v = float(raw.replace(",", ".").replace("%", ""))
            discount_rate = v / 100 if v > 1 else v
            break
        except ValueError:
            print("  ✗ Введіть число, наприклад: 8")

    # Cap Rate
    while True:
        raw = input(
            f"  Ставка капіталізації Cap Rate (%) "
            f"[орієнтир {default_cap*100:.0f}%, Enter = прийняти]: "
        ).strip()
        if not raw:
            cap_rate = default_cap
            break
        try:
            v = float(raw.replace(",", ".").replace("%", ""))
            cap_rate = v / 100 if v > 1 else v
            break
        except ValueError:
            print("  ✗ Введіть число, наприклад: 11")

    print(f"\n  → Знижка на торг : {discount_rate*100:.1f}%")
    print(f"  → Cap Rate        : {cap_rate*100:.1f}%")
    print(f"{'─'*52}\n")
    return discount_rate, cap_rate


# ─────────────────────────────────────────────────────────────────────────────
# Завантаження аналогів з Clippings/*.md
# ─────────────────────────────────────────────────────────────────────────────

def load_analogs_from_clippings(obj_type: str, deal_type: str) -> list[dict]:
    """Шукає розпарсені кліпінги відповідної категорії та типу угоди."""
    results = []
    for md_path in sorted(CLIPPINGS_DIR.glob("**/*.md")):
        fm = load_yaml_from_md(md_path)
        if not fm.get("parsed"):
            continue
        if fm.get("Category") != obj_type or fm.get("Deal_Type") != deal_type:
            continue

        analog: dict[str, Any] = {
            "name":      md_path.stem[:50],
            "area":      fm.get("Area"),
            "zone":      fm.get("Location_Zone", "Periphery"),
            "condition": fm.get("Condition_Type", "з ремонтом"),
        }
        if deal_type == "Sale":
            analog["price_psm"] = fm.get("Price_per_sqm")
        else:
            analog["price_psm"] = fm.get("Rent_per_sqm")

        analog["ceiling"]   = fm.get("Ceiling_Height")
        analog["power_kw"]  = fm.get("Power_kW")
        analog["railway"]   = fm.get("Railway", False)
        analog["class"]     = fm.get("Building_Class")
        analog["metro_min"] = fm.get("Distance_to_Metro")

        if analog.get("price_psm"):
            results.append(analog)
    return results


# ─────────────────────────────────────────────────────────────────────────────
# [B1] Порівняльний підхід
# ─────────────────────────────────────────────────────────────────────────────

def _area_adj(subject_area: float, analog_area: float | None) -> float:
    if not analog_area or analog_area <= 0:
        return 0.0
    ratio = subject_area / analog_area
    if ratio > 1.5:  return +0.06
    if ratio > 1.25: return +0.03
    if ratio < 0.67: return -0.06
    if ratio < 0.80: return -0.03
    return 0.0


def _zone_adj(subject_zone: str, analog_zone: str) -> float:
    diff = ZONE_RANK.get(subject_zone, 2) - ZONE_RANK.get(analog_zone, 2)
    return diff * 0.08


def _condition_adj(analog_condition: str | None) -> float:
    c = (analog_condition or "").lower()
    if "після будівельників" in c or "без ремонт" in c: return +0.15
    if "під оздоблення" in c:                            return +0.10
    if "з меблями" in c or "дизайнер" in c:             return -0.05
    return 0.0


def _tech_adj_warehouse(subject: dict, analog: dict) -> float:
    adj = 0.0
    s_ceil = subject.get("ceiling_height", 6.0)
    a_ceil = analog.get("ceiling")
    if a_ceil:
        if a_ceil >= 10 and s_ceil < 10:  adj -= 0.08
        elif a_ceil >= 8 and s_ceil < 8:  adj -= 0.05
        elif a_ceil < 5  and s_ceil >= 5: adj += 0.05
    if subject.get("railway") and not analog.get("railway"):     adj += 0.05
    if not subject.get("railway") and analog.get("railway"):     adj -= 0.05
    return adj


def _tech_adj_office(subject: dict, analog: dict) -> float:
    class_rank = {"A": 4, "B+": 3, "B": 2, "C": 1}
    s_cls = class_rank.get(subject.get("building_class", "B"), 2)
    a_cls = class_rank.get(analog.get("class"), 2)
    adj   = (s_cls - a_cls) * 0.05
    s_m   = subject.get("metro_min") or 10
    a_m   = analog.get("metro_min") or 15
    if s_m <= 3 and a_m > 10: adj += 0.05
    if s_m > 10 and a_m <= 3: adj -= 0.05
    return adj


def comparative_approach(subject: dict, analogs: list[dict],
                         discount_rate: float | None = None) -> dict:
    """[B1] Порівняльний підхід — матриця корегувань."""
    rows     = []
    obj_type = subject.get("type", "Warehouse")
    disc     = discount_rate if discount_rate is not None else \
               DISCOUNT_RATES.get(obj_type, 0.07)

    for a in analogs:
        psm = a.get("price_psm") or 0
        if not psm:
            continue
        psm_adj       = psm * (1 - disc)
        adj_scale     = _area_adj(subject["area"], a.get("area"))
        adj_location  = _zone_adj(subject.get("zone", "Periphery"),
                                  a.get("zone", "Periphery"))
        adj_condition = _condition_adj(a.get("condition"))
        if obj_type == "Warehouse":
            adj_tech = _tech_adj_warehouse(subject, a)
        elif obj_type == "Office":
            adj_tech = _tech_adj_office(subject, a)
        else:
            adj_tech = 0.0
        total_adj = adj_scale + adj_location + adj_condition + adj_tech
        final_psm = psm_adj * (1 + total_adj)
        rows.append({
            "name":          a.get("name", "—"),
            "district":      a.get("district", "—"),
            "area":          a.get("area"),
            "price_psm":     psm,
            "discount_adj":  round(psm_adj, 2),
            "adj_scale":     adj_scale,
            "adj_location":  adj_location,
            "adj_condition": adj_condition,
            "adj_tech":      adj_tech,
            "total_adj":     total_adj,
            "final_psm":     final_psm,
        })

    if not rows:
        return {"error": "Немає аналогів для порівняльного підходу"}

    final_psms = [r["final_psm"] for r in rows]
    avg_psm    = sum(final_psms) / len(final_psms)
    return {
        "rows":          rows,
        "avg_psm":       avg_psm,
        "value":         avg_psm * subject["area"],
        "n":             len(rows),
        "discount_rate": disc,
    }


# ─────────────────────────────────────────────────────────────────────────────
# [B2] Дохідний підхід
# ─────────────────────────────────────────────────────────────────────────────

def income_approach(subject: dict, cap_rate: float | None = None) -> dict:
    """[B2] Дохідний підхід: PGI → EGI → NOI → Капіталізація."""
    area      = subject["area"]
    rent_rate = subject.get("rent_rate", 6.0)
    vacancy   = subject.get("vacancy", 0.15)
    obj_type  = subject.get("type", "Warehouse")
    cap       = cap_rate if cap_rate is not None else \
                (subject.get("cap_rate") or CAP_RATES.get(obj_type, 0.11))
    opex_pct  = subject.get("opex_pct") or OPEX_PCT.get(obj_type, 0.18)

    pgi          = area * rent_rate * 12
    vacancy_loss = pgi * vacancy
    egi          = pgi - vacancy_loss
    opex         = egi * opex_pct
    noi          = egi - opex
    value_b2     = noi / cap

    return {
        "area": area, "rent_rate": rent_rate, "vacancy": vacancy,
        "cap_rate": cap, "opex_pct": opex_pct,
        "pgi": pgi, "vacancy_loss": vacancy_loss,
        "egi": egi, "opex": opex, "noi": noi, "value": value_b2,
    }


# ─────────────────────────────────────────────────────────────────────────────
# [B3] Узгодження результатів
# ─────────────────────────────────────────────────────────────────────────────

def reconciliation(b1: dict, b2: dict,
                   weight_b1: float = 0.50,
                   weight_b2: float = 0.50) -> dict:
    """[B3] Узгодження вартості двох підходів."""
    v1, v2 = b1.get("value", 0), b2.get("value", 0)
    if v1 and not v2:
        return {"value": v1, "weight_b1": 1.0, "weight_b2": 0.0,
                "value_b1": v1, "value_b2": 0, "note": "Тільки B1"}
    if v2 and not v1:
        return {"value": v2, "weight_b1": 0.0, "weight_b2": 1.0,
                "value_b1": 0, "value_b2": v2, "note": "Тільки B2"}
    weighted  = v1 * weight_b1 + v2 * weight_b2
    magnitude = 10 ** (len(str(int(weighted))) - 2)
    rounded   = round(weighted / magnitude) * magnitude
    return {
        "value_b1": v1, "value_b2": v2,
        "weight_b1": weight_b1, "weight_b2": weight_b2,
        "weighted": weighted, "value": rounded,
        "note": (f"Зважена ({weight_b1*100:.0f}% B1 + {weight_b2*100:.0f}% B2), "
                 f"округлено до {magnitude:,}"),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Стилі Excel
# ─────────────────────────────────────────────────────────────────────────────

_C_DARK   = "1C2E44"
_C_MID    = "2D4A6B"
_C_ACCENT = "D5B58A"
_C_LIGHT  = "EAF0F6"
_C_INPUT  = "E1FAFF"   # ← вхідні дані (редаговані)
_C_WHITE  = "FFFFFF"
_C_GREEN  = "D5E8D4"
_C_RED    = "F8CECC"
_C_YELL   = "FFF2CC"
_C_STRIPE = "F2F6FA"

_thin       = Side(style="thin",   color="BBBBBB")
_thick      = Side(style="medium", color=_C_DARK)
_border_all = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)


def _style(cell: Any, bold=False, italic=False, sz=10,
           color=_C_DARK, bg=_C_WHITE, align="center",
           wrap=False, fmt: str | None = None) -> None:
    """Застосовує форматування до комірки."""
    cell.font      = Font(name="Calibri", bold=bold, italic=italic,
                          size=sz, color=color)
    cell.fill      = PatternFill("solid", fgColor=bg)
    cell.alignment = Alignment(horizontal=align, vertical="center",
                               wrap_text=wrap)
    cell.border    = _border_all
    if fmt:
        cell.number_format = fmt


def _write(ws: Any, row: int, col: int, value: Any, **kw) -> Any:
    """Записує значення або формулу і стилізує комірку."""
    cell = ws.cell(row=row, column=col, value=value)
    _style(cell, **kw)
    return cell


def _hdr(ws: Any, row: int, col: int, text: str,
         width: float | None = None, dark: bool = True, sz: int = 10) -> None:
    """Заголовочна комірка (темний або золотий фон)."""
    bg = _C_DARK if dark else _C_ACCENT
    fc = _C_WHITE if dark else _C_DARK
    _write(ws, row, col, text, bold=True, sz=sz, color=fc,
           bg=bg, align="center", wrap=True)
    if width:
        ws.column_dimensions[get_column_letter(col)].width = width


def _merge_write(ws: Any, row: int, c1: int, c2: int, text: str,
                 bold=True, sz=10, bg=_C_DARK, color=_C_WHITE,
                 align="left") -> None:
    """Записує і об'єднує комірки."""
    ws.merge_cells(start_row=row, start_column=c1,
                   end_row=row,   end_column=c2)
    cell = ws.cell(row=row, column=c1, value=text)
    _style(cell, bold=bold, sz=sz, color=color, bg=bg, align=align)


def _input_cell(ws: Any, row: int, col: int, value: Any,
                fmt: str = "General") -> Any:
    """Вхідна комірка з підсвіченням E1FAFF."""
    cell = ws.cell(row=row, column=col, value=value)
    _style(cell, bold=True, color=_C_DARK, bg=_C_INPUT,
           align="center", fmt=fmt)
    return cell


def _pct(v: float) -> str:
    return f"{v:+.1%}" if v != 0 else "0%"


# ─────────────────────────────────────────────────────────────────────────────
# Генерація Excel [B1 + B2 + B3] з Excel-формулами
# ─────────────────────────────────────────────────────────────────────────────

def generate_excel(subject: dict, b1: dict, b2: dict, b3: dict,
                   out_path: Path,
                   discount_rate: float, cap_rate: float) -> None:
    """
    Створює Excel з трьома аркушами.

    Ключові принципи:
      • Вхідні дані — підсвічені E1FAFF, редаговані користувачем
      • Всі проміжні розрахунки — Excel-формули (не hardcode)
      • B3 посилається на B1 і B2 через cross-sheet формули
    """
    wb = openpyxl.Workbook()

    # ════════════════════════════════════════════════════════════════════════
    # АРКУШ 1: B1_Порівняльний
    # ════════════════════════════════════════════════════════════════════════
    ws1 = wb.active
    ws1.title = "B1_Порівняльний"
    ws1.sheet_view.showGridLines = False

    ws1.row_dimensions[1].height = 8
    ws1.row_dimensions[2].height = 32
    _merge_write(ws1, 2, 1, 12,
                 f"[B1]  ПОРІВНЯЛЬНИЙ ПІДХІД — МАТРИЦЯ КОРЕГУВАНЬ  "
                 f"|  {subject.get('name', 'Об єкт оцінки')}",
                 sz=12, align="center")

    # ── Вхідний параметр: знижка на торг ──
    ws1.row_dimensions[4].height = 22
    _write(ws1, 4, 1, "Знижка на торг, %  ⚙ (E1FAFF — редагована)",
           bold=True, bg=_C_LIGHT, align="left")
    ws1.merge_cells("A4:E4")
    DISC_CELL = "F4"   # адреса вхідного параметра
    _input_cell(ws1, 4, 6, discount_rate, fmt="0.0%")
    _write(ws1, 4, 7, "← змінюйте для перерахунку всієї таблиці",
           italic=True, color="777777", align="left", sz=9)
    ws1.merge_cells("G4:L4")
    ws1.column_dimensions["F"].width = 10

    # ── Заголовки таблиці аналогів ──
    HDR_ROW = 6
    ws1.row_dimensions[5].height = 6
    ws1.row_dimensions[HDR_ROW].height = 42

    col_defs = [
        (1,  "№",                      4),
        (2,  "Аналог",                22),
        (3,  "Район",                 14),
        (4,  "Площа,\nм²",             8),
        (5,  "Ціна пропозиції\n$/м²", 12),
        (6,  "Після торгу\n$/м²",     12),
        (7,  "Кор.\nмасштаб",          9),
        (8,  "Кор.\nлокація",          9),
        (9,  "Кор.\nстан",             9),
        (10, "Кор.\nтехн.",            9),
        (11, "Σ\nкорег.",              9),
        (12, "Скоригована\nціна $/м²",13),
    ]
    for col_n, label, width in col_defs:
        _hdr(ws1, HDR_ROW, col_n, label, width)

    # ── Рядки аналогів ──
    rows_b1 = b1.get("rows", [])
    DATA_START = HDR_ROW + 1

    for i, row in enumerate(rows_b1):
        r  = DATA_START + i
        bg = _C_WHITE if i % 2 == 0 else _C_STRIPE
        ws1.row_dimensions[r].height = 18

        _write(ws1, r, 1, i + 1,                   bg=bg, fmt="0")
        _write(ws1, r, 2, row["name"],              bg=bg, align="left")
        _write(ws1, r, 3, row.get("district", "—"), bg=bg, align="left")

        # Площа та ціна — E1FAFF (вхідні, верифіковані аналітиком)
        _input_cell(ws1, r, 4, row.get("area"),      fmt="#,##0")
        _input_cell(ws1, r, 5, row.get("price_psm"), fmt='#,##0.00 "$"')

        # Col F = Ціна після торгу = ФОРМУЛА
        c_e = get_column_letter(5)
        c_f = get_column_letter(6)
        cell_f = ws1.cell(row=r, column=6,
                          value=f"={c_e}{r}*(1-{DISC_CELL})")
        _style(cell_f, bold=True, bg=bg, fmt='#,##0.00 "$"')

        # Корегування (обчислені Python, підсвічені за знаком)
        adj_cols = [
            (7,  row["adj_scale"]),
            (8,  row["adj_location"]),
            (9,  row["adj_condition"]),
            (10, row["adj_tech"]),
        ]
        for c_n, val in adj_cols:
            clr = _C_GREEN if val > 0 else (_C_RED if val < 0 else _C_YELL)
            _write(ws1, r, c_n, val, bg=clr, fmt="0.0%")

        # Col K = Σ корегувань = ФОРМУЛА
        cg = get_column_letter(7); ch = get_column_letter(8)
        ci = get_column_letter(9); cj = get_column_letter(10)
        cell_sum = ws1.cell(row=r, column=11,
                            value=f"={cg}{r}+{ch}{r}+{ci}{r}+{cj}{r}")
        _style(cell_sum, bold=True, fmt="0.0%",
               bg=(_C_YELL if abs(row["total_adj"]) > 0.15 else _C_GREEN))

        # Col L = Скоригована ціна = ФОРМУЛА
        ck = get_column_letter(11)
        cl = get_column_letter(12)
        cell_fin = ws1.cell(row=r, column=12,
                            value=f"={c_f}{r}*(1+{ck}{r})")
        _style(cell_fin, bold=True, bg=bg, fmt='#,##0.00 "$"')

    LAST_DATA = DATA_START + len(rows_b1) - 1
    lc        = get_column_letter(12)

    # ── Статистика (Excel-формули) ──
    STAT_HDR = LAST_DATA + 2
    ws1.row_dimensions[LAST_DATA + 1].height = 6
    ws1.row_dimensions[STAT_HDR].height = 24
    _merge_write(ws1, STAT_HDR, 1, 11,
                 "СТАТИСТИКА СКОРИГОВАНИХ ЦІН, $/м²", sz=10, bg=_C_MID)

    stats = [
        ("Мін.",             f"=MIN({lc}{DATA_START}:{lc}{LAST_DATA})"),
        ("Середня (AVERAGE)", f"=AVERAGE({lc}{DATA_START}:{lc}{LAST_DATA})"),
        ("Медіана",           f"=MEDIAN({lc}{DATA_START}:{lc}{LAST_DATA})"),
        ("Макс.",             f"=MAX({lc}{DATA_START}:{lc}{LAST_DATA})"),
    ]
    AVG_ROW = STAT_HDR + 2   # рядок "Середня"
    for k, (lbl, formula) in enumerate(stats):
        sr = STAT_HDR + 1 + k
        ws1.row_dimensions[sr].height = 17
        _write(ws1, sr, 1, lbl, bold=True, bg=_C_STRIPE, align="left")
        ws1.merge_cells(start_row=sr, start_column=1, end_row=sr, end_column=11)
        cell_s = ws1.cell(row=sr, column=12, value=formula)
        _style(cell_s, bold=True, bg=_C_LIGHT, fmt='#,##0.00 "$"')

    # ── Прийнята ціна (= AVERAGE) ──
    ACC_ROW = STAT_HDR + len(stats) + 2
    ws1.row_dimensions[ACC_ROW - 1].height = 6
    ws1.row_dimensions[ACC_ROW].height = 24
    _merge_write(ws1, ACC_ROW, 1, 11,
                 "✅  ПРИЙНЯТА РИНКОВА ЦІНА ($/м²) — AVERAGE",
                 sz=10, bg=_C_ACCENT, color=_C_DARK)
    AVG_CELL = f"{lc}{AVG_ROW}"
    ACC_CELL = f"{lc}{ACC_ROW}"
    cell_acc = ws1.cell(row=ACC_ROW, column=12, value=f"={AVG_CELL}")
    _style(cell_acc, bold=True, sz=11, bg=_C_ACCENT, color=_C_DARK,
           fmt='#,##0.00 "$"')

    # ── Вартість B1 ──
    VAL_B1_ROW  = ACC_ROW + 1
    B1_VAL_CELL = f"{lc}{VAL_B1_ROW}"   # для cross-sheet посилання з B3
    ws1.row_dimensions[VAL_B1_ROW].height = 26
    _merge_write(ws1, VAL_B1_ROW, 1, 11,
                 f"ВАРТІСТЬ (B1) = $/м² × {subject['area']:,.0f} м²",
                 sz=11, bg=_C_DARK)
    cell_v1 = ws1.cell(row=VAL_B1_ROW, column=12,
                       value=f"={ACC_CELL}*{subject['area']}")
    _style(cell_v1, bold=True, sz=11, bg=_C_DARK, color=_C_WHITE,
           fmt='#,##0 "$"')

    # ════════════════════════════════════════════════════════════════════════
    # АРКУШ 2: B2_Дохідний
    # ════════════════════════════════════════════════════════════════════════
    ws2 = wb.create_sheet("B2_Дохідний")
    ws2.sheet_view.showGridLines = False
    for col_letter, width in [("A", 38), ("B", 22), ("C", 16), ("D", 14)]:
        ws2.column_dimensions[col_letter].width = width

    ws2.row_dimensions[1].height = 8
    ws2.row_dimensions[2].height = 32
    _merge_write(ws2, 2, 1, 3,
                 "[B2]  ДОХІДНИЙ ПІДХІД  |  Пряма капіталізація NOI",
                 sz=12, align="center")

    r2 = 4
    ws2.row_dimensions[r2].height = 24
    _merge_write(ws2, r2, 1, 3,
                 "ВХІДНІ ПАРАМЕТРИ  (комірки E1FAFF — редаговані)", sz=10, bg=_C_MID)
    r2 += 1

    def _b2_input(label: str, val: Any, fmt: str, note: str = "") -> str:
        nonlocal r2
        ws2.row_dimensions[r2].height = 18
        _write(ws2, r2, 1, label, bold=True, bg=_C_STRIPE, align="left")
        _input_cell(ws2, r2, 2, val, fmt=fmt)
        addr = f"B{r2}"
        if note:
            _write(ws2, r2, 3, note, italic=True, color="777777",
                   align="left", sz=9)
            ws2.merge_cells(f"C{r2}:D{r2}")
        r2 += 1
        return addr

    area_a    = _b2_input("GLA — площа, м²",                    b2["area"],     "#,##0",
                           "GBA з урахуванням усіх поверхів")
    rent_a    = _b2_input("Орендна ставка, $/м²/міс",           b2["rent_rate"], '0.00 "$"',
                           "Ринкова (з B1 або ринок)")
    vac_a     = _b2_input("Вакансія та втрати від несплати",     b2["vacancy"],   "0%",
                           "Ринковий орієнтир + надбавка")
    opex_a    = _b2_input("OPEX від EGI, %",                    b2["opex_pct"],  "0%",
                           "Охорона, ЖКГ, страхування, адмін.")
    caprate_a = _b2_input("Cap Rate (ставка капіт.), %",        cap_rate,        "0.0%",
                           f"Ринк. UA 2024-25; введено: {cap_rate*100:.1f}%")
    r2 += 1

    ws2.row_dimensions[r2].height = 24
    _merge_write(ws2, r2, 1, 3,
                 "РОЗРАХУНОК NOI  (всі комірки — формули Excel)", sz=10, bg=_C_MID)
    r2 += 1

    def _b2_formula(label: str, formula: str, fmt: str,
                    bold=False, bg=_C_WHITE, note="") -> str:
        nonlocal r2
        ws2.row_dimensions[r2].height = 18
        _write(ws2, r2, 1, label, bold=bold, bg=bg, align="left")
        cell_f = ws2.cell(row=r2, column=2, value=formula)
        _style(cell_f, bold=bold, bg=bg, fmt=fmt)
        addr = f"B{r2}"
        if note:
            _write(ws2, r2, 3, note, italic=True, color="777777",
                   align="left", sz=9)
            ws2.merge_cells(f"C{r2}:D{r2}")
        r2 += 1
        return addr

    pgi_a    = _b2_formula("PGI  — Потенційний валовий дохід, $/рік",
                            f"={area_a}*{rent_a}*12",       '#,##0 "$"',
                            bold=True, note=f"={area_a} × {rent_a} × 12")
    vacloss_a = _b2_formula("  − Втрати від вакансії та несплати, $/рік",
                             f"=-{pgi_a}*{vac_a}",          '#,##0 "$"',
                             note=f"PGI × {vac_a}")
    egi_a    = _b2_formula("EGI  — Ефективний валовий дохід, $/рік",
                            f"={pgi_a}+{vacloss_a}",        '#,##0 "$"',
                            bold=True, bg=_C_LIGHT)
    opex_v_a = _b2_formula("  − OPEX (операційні витрати), $/рік",
                            f"=-{egi_a}*{opex_a}",          '#,##0 "$"')
    noi_a    = _b2_formula("NOI  — Чистий операційний дохід, $/рік",
                            f"={egi_a}+{opex_v_a}",         '#,##0 "$"',
                            bold=True, bg=_C_LIGHT)
    r2 += 1

    B2_VAL_ROW = r2
    b2_val_a   = _b2_formula("ВАРТІСТЬ (B2) = NOI / Cap Rate",
                              f"={noi_a}/{caprate_a}",      '#,##0 "$"',
                              bold=True, bg=_C_ACCENT)
    ws2.cell(row=B2_VAL_ROW, column=1).font = Font(
        name="Calibri", bold=True, size=12, color=_C_DARK)
    ws2.cell(row=B2_VAL_ROW, column=2).font = Font(
        name="Calibri", bold=True, size=12, color=_C_DARK)

    # ════════════════════════════════════════════════════════════════════════
    # АРКУШ 3: B3_Узгодження
    # ════════════════════════════════════════════════════════════════════════
    ws3 = wb.create_sheet("B3_Узгодження")
    ws3.sheet_view.showGridLines = False
    for col_letter, width in [("A", 42), ("B", 24), ("C", 16)]:
        ws3.column_dimensions[col_letter].width = width

    ws3.row_dimensions[1].height = 8
    ws3.row_dimensions[2].height = 32
    _merge_write(ws3, 2, 1, 3,
                 "[B3]  УЗГОДЖЕННЯ РЕЗУЛЬТАТІВ  |  Зважена вартість",
                 sz=12, align="center")

    r3 = 4
    ws3.row_dimensions[r3].height = 24
    _merge_write(ws3, r3, 1, 3,
                 "РЕЗУЛЬТАТИ ПІДХОДІВ (cross-sheet формули)", sz=10, bg=_C_MID)
    r3 += 1

    def _b3_row(label: str, value: Any, fmt: str, bold=False,
                bg=_C_WHITE, note: str = "") -> str:
        nonlocal r3
        ws3.row_dimensions[r3].height = 18
        _write(ws3, r3, 1, label, bold=bold, bg=bg, align="left")
        cell_v = ws3.cell(row=r3, column=2, value=value)
        _style(cell_v, bold=bold, bg=bg, fmt=fmt)
        if note:
            _write(ws3, r3, 3, note, italic=True, color="777777",
                   align="left", sz=9)
        addr = f"B{r3}"
        r3 += 1
        return addr

    # Cross-sheet посилання: одинарні лапки обов'язкові для кирилиці
    b1_ref_addr = _b3_row(
        "Вартість за порівняльним підходом (B1), $",
        f"='B1_Порівняльний'!{B1_VAL_CELL}",
        '#,##0 "$"', note="← B1_Порівняльний")

    b2_ref_addr = _b3_row(
        "Вартість за дохідним підходом (B2), $",
        f"='B2_Дохідний'!{b2_val_a}",
        '#,##0 "$"', note="← B2_Дохідний")

    r3 += 1
    ws3.row_dimensions[r3].height = 24
    _merge_write(ws3, r3, 1, 3,
                 "ВАГИ ПІДХОДІВ  (E1FAFF — редаговані)", sz=10, bg=_C_MID)
    r3 += 1

    ws3.row_dimensions[r3].height = 18
    _write(ws3, r3, 1, "Вага B1 (порівняльний підхід)",
           bold=True, bg=_C_STRIPE, align="left")
    W1_CELL = f"B{r3}"
    _input_cell(ws3, r3, 2, b3.get("weight_b1", 0.50), fmt="0%")
    r3 += 1

    ws3.row_dimensions[r3].height = 18
    _write(ws3, r3, 1, "Вага B2 (дохідний підхід)",
           bold=True, bg=_C_STRIPE, align="left")
    W2_CELL = f"B{r3}"
    _input_cell(ws3, r3, 2, b3.get("weight_b2", 0.50), fmt="0%")
    r3 += 1

    r3 += 1
    ws3.row_dimensions[r3].height = 24
    _merge_write(ws3, r3, 1, 3, "УЗГОДЖЕНА ВАРТІСТЬ", sz=10, bg=_C_MID)
    r3 += 1

    wtd_addr = _b3_row(
        "Зважена вартість (до округлення), $",
        f"={b1_ref_addr}*{W1_CELL}+{b2_ref_addr}*{W2_CELL}",
        '#,##0 "$"', bold=True, bg=_C_LIGHT)

    r3 += 1
    ws3.row_dimensions[r3].height = 32
    _merge_write(ws3, r3, 1, 1,
                 "ПІДСУМКОВА ВАРТІСТЬ (узгоджена), $",
                 sz=12, bg=_C_DARK, bold=True)
    ws3.cell(row=r3, column=1).font = Font(
        name="Calibri", bold=True, size=12, color=_C_WHITE)
    cell_final = ws3.cell(row=r3, column=2,
                          value=f"=ROUND({wtd_addr},-4)")
    _style(cell_final, bold=True, sz=13, bg=_C_ACCENT, color=_C_DARK,
           fmt='#,##0 "$"')

    r3 += 2
    ws3.row_dimensions[r3].height = 36
    _merge_write(ws3, r3, 1, 3,
                 f"⚠ Попередня оцінка в аналітичних цілях. "
                 f"Дата: {date.today().strftime('%d.%m.%Y')}. "
                 f"Методологія: [B1]+[B2]+[B3].",
                 bold=False, sz=9, bg=_C_YELL, color=_C_DARK, align="left")

    wb.save(out_path)
    print(f"  ✓ Excel: {out_path.name}")


# ─────────────────────────────────────────────────────────────────────────────
# Генерація Word
# ─────────────────────────────────────────────────────────────────────────────

def _docx_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1C, 0x2E, 0x44)


def _docx_table_2col(doc: Document, rows: list[tuple[str, str]],
                     col_widths: tuple[float, float] = (9.0, 7.5)) -> None:
    """Двоколонкова таблиця (параметр | значення)."""
    tbl = doc.add_table(rows=len(rows) + 1, cols=2)
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    hdr_cells[0].text = "Параметр"
    hdr_cells[1].text = "Значення"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (label, val) in enumerate(rows, start=1):
        tbl.cell(i, 0).text = label
        tbl.cell(i, 1).text = str(val)
    for row in tbl.rows:
        for j, w in enumerate(col_widths):
            row.cells[j].width = Cm(w)


def _docx_analogs_table(doc: Document, rows_b1: list[dict]) -> None:
    """Таблиця аналогів [B1] у Word."""
    cols = ["Аналог", "Площа,\nм²", "Ціна,\n$/м²",
            "Після\nторгу", "Σ кор-нь", "Підсумк.\nціна, $/м²"]
    tbl = doc.add_table(rows=len(rows_b1) + 1, cols=len(cols))
    tbl.style = "Table Grid"
    for j, h in enumerate(cols):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for i, row in enumerate(rows_b1, start=1):
        vals = [
            row["name"],
            f"{row.get('area') or '—':,}" if row.get("area") else "—",
            f"{row.get('price_psm') or '—':,.0f}" if row.get("price_psm") else "—",
            f"{row.get('discount_adj') or '—':,.0f}" if row.get("discount_adj") else "—",
            _pct(row["total_adj"]),
            f"{row['final_psm']:,.0f}",
        ]
        for j, v in enumerate(vals):
            tbl.cell(i, j).text = str(v)


def generate_word(subject: dict, b1: dict, b2: dict, b3: dict,
                  out_path: Path) -> None:
    """Генерує Word-звіт [B1, B2, B3] українською мовою."""
    doc = Document()

    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.0)

    obj_name  = subject.get("name", "Об'єкт")
    obj_type  = subject.get("type", "Warehouse")
    obj_area  = subject["area"]
    obj_zone  = subject.get("zone", "—")
    today_str = date.today().strftime("%d.%m.%Y")

    # Титульний блок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"ЗВІТ ПРО ОЦІНКУ РИНКОВОЇ ВАРТОСТІ\n{obj_name.upper()}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1C, 0x2E, 0x44)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Станом на {today_str}  ·  Тип: {obj_type}  ·  Площа: {obj_area:,.0f} м²")
    doc.add_paragraph()

    # 1. Резюме
    _docx_heading(doc, "1. Резюме (Executive Summary)", level=1)
    doc.add_paragraph(
        f"Об'єктом оцінки є {obj_type.lower()} загальною площею {obj_area:,.0f} м², "
        f"розташований у зоні «{obj_zone}». Оцінка виконана за підходами [B1] та [B2], "
        f"результати узгоджено за методологією [B3]."
    )
    summary_rows = [
        ("Підсумкова ринкова вартість",        f"$ {b3['value']:,.0f}"),
        ("  — порівняльний підхід (B1)",       f"$ {b3['value_b1']:,.0f}"),
        ("  — дохідний підхід (B2)",           f"$ {b3['value_b2']:,.0f}"),
        ("Вартість за 1 м² (узгоджена)",       f"$ {b3['value'] / obj_area:,.0f}"),
        ("NOI (рік)",                          f"$ {b2['noi']:,.0f}"),
        ("Cap Rate (ринковий)",                f"{b2['cap_rate']*100:.1f}%"),
        ("Дата оцінки",                        today_str),
    ]
    _docx_table_2col(doc, summary_rows)
    doc.add_paragraph()

    # 2. Методологія
    _docx_heading(doc, "2. Методологія оцінки", level=1)
    for bullet in [
        "[B1] Порівняльний підхід — матриця кількісних корегувань (торг, масштаб, "
        "локація, стан, техпараметри).",
        "[B2] Дохідний підхід — пряма капіталізація NOI. PGI → EGI (вакансія) → NOI (OPEX).",
        "[B3] Узгодження — зважена середня з урахуванням якості бази аналогів.",
    ]:
        p = doc.add_paragraph(bullet, style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)
    doc.add_paragraph()

    # 3. Порівняльний підхід [B1]
    _docx_heading(doc, "3. Порівняльний підхід [B1]", level=1)
    doc.add_paragraph(
        f"Відібрано {b1['n']} ринкових аналогів. Знижка на торг: "
        f"{b1.get('discount_rate', 0)*100:.1f}%."
    )
    _docx_analogs_table(doc, b1.get("rows", []))
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Результат [B1]: ").bold = True
    p.add_run(f"середня скоригована ціна — ${b1.get('avg_psm', 0):,.0f}/м²; "
              f"вартість = ${b1.get('value', 0):,.0f}")
    doc.add_paragraph()

    # 4. Дохідний підхід [B2]
    _docx_heading(doc, "4. Дохідний підхід [B2]", level=1)
    income_rows = [
        ("GLA, м²",                    f"{b2['area']:,.0f}"),
        ("Орендна ставка, $/м²/міс",   f"{b2['rent_rate']:.2f}"),
        ("PGI, $/рік",                 f"{b2['pgi']:,.0f}"),
        ("Вакансія та несплата, $/рік",f"({b2['vacancy_loss']:,.0f})"),
        ("EGI, $/рік",                 f"{b2['egi']:,.0f}"),
        ("OPEX, $/рік",                f"({b2['opex']:,.0f})"),
        ("NOI, $/рік",                 f"{b2['noi']:,.0f}"),
        ("Cap Rate",                   f"{b2['cap_rate']*100:.1f}%"),
        ("Вартість (B2)",              f"$ {b2['value']:,.0f}"),
    ]
    _docx_table_2col(doc, income_rows)
    doc.add_paragraph()

    # 5. Узгодження [B3]
    _docx_heading(doc, "5. Узгодження результатів [B3]", level=1)
    rec_rows = [
        ("Вартість за B1, $",          f"{b3['value_b1']:,.0f}"),
        ("Вага B1",                    f"{b3['weight_b1']*100:.0f}%"),
        ("Вартість за B2, $",          f"{b3['value_b2']:,.0f}"),
        ("Вага B2",                    f"{b3['weight_b2']*100:.0f}%"),
        ("Зважена (до округлення), $", f"{b3.get('weighted', b3['value']):,.0f}"),
        ("ПІДСУМКОВА ВАРТІСТЬ, $",     f"{b3['value']:,.0f}"),
    ]
    _docx_table_2col(doc, rec_rows)
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Примітка: ").bold = True
    note.add_run(b3.get("note", ""))
    doc.add_paragraph()

    # 6. Ризики
    _docx_heading(doc, "6. Ризики та обмеження", level=1)
    for risk in [
        "Воєнний стан: ринок функціонує з обмеженим попитом та зниженою ліквідністю.",
        "Обмежена база аналогів: частина з відкритих пропозицій (OLX), що може "
        "відображати завищені ціни.",
        "Курс валют: зміна UAH/USD впливає на гривневий еквівалент.",
        "Юридичні ризики: обтяження або незавершені судові справи можуть знизити вартість.",
    ]:
        p = doc.add_paragraph(risk, style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)
    doc.add_paragraph()

    # 7. Висновок
    _docx_heading(doc, "7. Висновок", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Ринкова вартість об'єкта «{obj_name}» "
              f"(тип: {obj_type}, площа: {obj_area:,.0f} м², зона: {obj_zone}) "
              f"станом на {today_str} складає:")
    final_p = doc.add_paragraph()
    final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final_p.add_run(f"$ {b3['value']:,.0f}  (USD)")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1C, 0x2E, 0x44)
    doc.add_paragraph()
    disc = doc.add_paragraph()
    disc.add_run("Застереження: ").bold = True
    disc.add_run(
        "Звіт підготовлено в інформаційних цілях і не є офіційним висновком "
        "суб'єкта оціночної діяльності. Для юридично значущих операцій "
        "необхідна сертифікована оцінка."
    )
    doc.save(out_path)
    print(f"  ✓ Word:  {out_path.name}")


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

def build_subject(args: argparse.Namespace) -> dict:
    """Збирає dict суб'єкта оцінки з аргументів CLI."""
    return {
        "name":              args.name,
        "type":              args.type,
        "area":              args.area,
        "zone":              args.zone,
        "address":           args.address,
        "condition":         args.condition,
        "asking_price":      args.price,
        "rent_rate":         args.rent_rate,
        "vacancy":           args.vacancy,
        "cap_rate":          args.cap_rate,
        "opex_pct":          args.opex_pct,
        # Warehouse extras
        "ceiling_height":    args.ceiling,
        "railway":           args.railway,
        # Office extras
        "building_class":    args.building_class,
        "metro_min":         args.metro_min,
        # Audit fields (None = будуть запитані при --audit)
        "district":          None,
        "floor_type":        None,
        "renovation_style":  None,
        "generator":         None,
        "shelter":           None,
        "parking_underground": None,
        "parking_open":      None,
        "opex_owner":        None,
    }


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Генератор звіту про оцінку КН [B1, B2, B3]  |  Full Data Audit"
    )
    parser.add_argument("--name",    default="Об'єкт оцінки", help="Назва об'єкта")
    parser.add_argument("--type",    default="Warehouse",
                        choices=["Warehouse", "Office", "Retail", "Land", "Residential"])
    parser.add_argument("--area",    type=float, required=True, help="Площа GBA, м²")
    parser.add_argument("--price",   type=float, default=None,  help="Ціна пропозиції, $")
    parser.add_argument("--zone",    default="Periphery",
                        choices=["Center", "Middle", "Periphery", "Suburbs"])
    parser.add_argument("--address", default="—")
    parser.add_argument("--condition", default="з ремонтом")
    # B2
    parser.add_argument("--rent-rate",  type=float, default=None)
    parser.add_argument("--vacancy",    type=float, default=0.15)
    parser.add_argument("--cap-rate",   type=float, default=None)
    parser.add_argument("--opex-pct",   type=float, default=None)
    # B3
    parser.add_argument("--weight-b1",  type=float, default=0.50)
    parser.add_argument("--weight-b2",  type=float, default=0.50)
    # Тип-специфічні
    parser.add_argument("--ceiling",        type=float, default=None)
    parser.add_argument("--railway",        action="store_true")
    parser.add_argument("--building-class", default=None,
                        choices=["A", "B+", "B", "C"])
    parser.add_argument("--metro-min",      type=int,   default=None)
    # Режими
    parser.add_argument("--no-audit",       action="store_true",
                        help="Вимкнути Full Data Audit")
    parser.add_argument("--no-interactive", action="store_true",
                        help="Не запитувати знижку і Cap Rate (використати defaults)")
    parser.add_argument("--object-dir",     default=None,
                        help="Ім'я папки об'єкта в Объекты/ (для FDA)")
    parser.add_argument("--out-dir",        default=str(OUTPUT_DIR))

    args = parser.parse_args()

    if abs(args.weight_b1 + args.weight_b2 - 1.0) > 0.01:
        sys.exit("Помилка: weight-b1 + weight-b2 мають дорівнювати 1.0")

    subject = build_subject(args)

    # ── Full Data Audit ───────────────────────────────────────────────────────
    md_path: Path | None = None
    if not args.no_audit:
        obj_dir_name = args.object_dir or args.name
        md_path = find_object_md(obj_dir_name)
        if md_path:
            print(f"  FDA: MD-картка знайдена → {md_path.name}")
        else:
            print(f"  FDA: MD-картка не знайдена для «{obj_dir_name}» "
                  f"(Объекты/{obj_dir_name}/wiki/objects/)")
        subject = full_data_audit(subject, obj_dir_name, md_path)

    # ── Динамічні вхідні параметри ────────────────────────────────────────────
    if not args.no_interactive:
        discount_rate, cap_rate = ask_dynamic_inputs(args.type)
    else:
        discount_rate = DISCOUNT_RATES.get(args.type, 0.07)
        cap_rate      = args.cap_rate or CAP_RATES.get(args.type, 0.11)
        print(f"  Знижка на торг: {discount_rate*100:.1f}%  "
              f"|  Cap Rate: {cap_rate*100:.1f}%  (defaults)")

    # ── Rent rate (якщо не вказано) ───────────────────────────────────────────
    if not subject.get("rent_rate"):
        defaults = {"Warehouse": 5.5, "Office": 14.0, "Retail": 18.0,
                    "Land": 3.0, "Residential": 10.0}
        subject["rent_rate"] = defaults.get(args.type, 6.0)
        print(f"  ⚠️  Припущення: rent_rate = {subject['rent_rate']} $/м²/міс")

    # ── Аналоги ───────────────────────────────────────────────────────────────
    clipping_analogs = load_analogs_from_clippings(args.type, "Sale")
    demo             = DEMO_ANALOGS.get(args.type, DEMO_ANALOGS["Warehouse"])
    if len(clipping_analogs) >= 5:
        analogs = clipping_analogs
        print(f"  Аналоги: {len(analogs)} кліпінгів")
    else:
        combined = clipping_analogs + [a for a in demo if a not in clipping_analogs]
        analogs  = combined[:max(len(combined), 5)]
        print(f"  Аналоги: {len(clipping_analogs)} кліпінгів "
              f"+ {len(analogs)-len(clipping_analogs)} демо")

    # ── Розрахунки ────────────────────────────────────────────────────────────
    print("\n[B1] Порівняльний підхід...")
    b1 = comparative_approach(subject, analogs, discount_rate)
    if "error" in b1:
        sys.exit(f"B1 помилка: {b1['error']}")
    print(f"     Аналогів: {b1['n']} | Середня: ${b1['avg_psm']:,.0f}/м² | "
          f"Вартість: ${b1['value']:,.0f}")

    print("[B2] Дохідний підхід...")
    b2 = income_approach(subject, cap_rate)
    print(f"     PGI: ${b2['pgi']:,.0f} | NOI: ${b2['noi']:,.0f} | "
          f"Cap {b2['cap_rate']*100:.0f}% | Вартість: ${b2['value']:,.0f}")

    print("[B3] Узгодження...")
    b3 = reconciliation(b1, b2, args.weight_b1, args.weight_b2)
    print(f"     Підсумкова вартість: ${b3['value']:,.0f}\n")

    # ── Генерація файлів ─────────────────────────────────────────────────────
    out_dir   = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^\w\-_а-яА-ЯіІїЇєЄ ]", "_", args.name
                       ).strip().replace(" ", "_")

    xlsx_path = out_dir / f"Report_Calc_{safe_name}.xlsx"
    docx_path = out_dir / f"Звіт_про_оцінку_{safe_name}.docx"

    generate_excel(subject, b1, b2, b3, xlsx_path, discount_rate, cap_rate)
    generate_word(subject, b1, b2, b3, docx_path)

    print(f"\n  Папка: {out_dir}")


if __name__ == "__main__":
    main()
