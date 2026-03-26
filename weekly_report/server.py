"""
Сервер для системы еженедельной отчётности.
Запуск: /usr/bin/python3 weekly_report/server.py
Доступ: http://localhost:8080
"""

import json
import os
import re
import hashlib
import datetime
from http.server import HTTPServer, BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs
from pathlib import Path

# ── Пути ──────────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent.parent
DATA_DIR = BASE_DIR / "data" / "reports"
HTML_FILE    = Path(__file__).parent / "app.html"
HTML_FILE_V2 = Path(__file__).parent / "app_v2.html"

DATA_DIR.mkdir(parents=True, exist_ok=True)

PROJECTS_FILE = DATA_DIR / "projects.json"


# ── Утиліти ────────────────────────────────────────────────────────────────────

def _load_json(path: Path, default):
    """Читає JSON-файл, повертає default якщо немає."""
    if path.exists():
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    return default


def _save_json(path: Path, data):
    """Зберігає дані у JSON-файл."""
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _week_file(iso_week: str) -> Path:
    """Повертає шлях до файлу тижня (напр. 2026-W12)."""
    return DATA_DIR / f"{iso_week}.json"


def _current_iso_week() -> str:
    """Повертає поточний ISO-тиждень у форматі 'YYYY-WXX'."""
    today = datetime.date.today()
    year, week, _ = today.isocalendar()
    return f"{year}-W{week:02d}"


def _week_bounds(iso_week: str) -> tuple[str, str]:
    """Повертає дати початку та кінця тижня за ISO-номером."""
    year, week = int(iso_week[:4]), int(iso_week[6:])
    monday = datetime.date.fromisocalendar(year, week, 1)
    friday = monday + datetime.timedelta(days=4)
    return monday.strftime("%Y-%m-%d"), friday.strftime("%Y-%m-%d")


def _format_date_ua(date_str: str) -> str:
    """Форматує дату 'YYYY-MM-DD' у 'DD.MM'."""
    d = datetime.date.fromisoformat(date_str)
    return d.strftime("%d.%m")


def _get_history(n: int) -> list:
    """
    Повертає зведення за останні n тижнів:
    [{iso_week, week_start, week_end, plan_count, done_count, pct, entries_by_project}]
    """
    result = []
    today = datetime.date.today()
    year, current_week, _ = today.isocalendar()

    for delta in range(n - 1, -1, -1):
        # Рахуємо тиждень назад від поточного
        target_monday = datetime.date.fromisocalendar(year, current_week, 1) - datetime.timedelta(weeks=delta)
        t_year, t_week, _ = target_monday.isocalendar()
        iso_week = f"{t_year}-W{t_week:02d}"
        week_start, week_end = _week_bounds(iso_week)

        data = _load_json(_week_file(iso_week), None)
        if data is None:
            result.append({
                "iso_week": iso_week,
                "week_start": week_start,
                "week_end": week_end,
                "plan_count": 0,
                "done_count": 0,
                "pct": 0,
                "has_data": False,
                "projects": {}
            })
            continue

        total_plan = 0
        total_done = 0
        projects_summary = {}

        for proj, entry in data.get("entries", {}).items():
            cancelled = set(entry.get("cancelled", []))
            plan = [t for t in entry.get("plan", []) if t not in cancelled]
            done = entry.get("done", [])
            extra = entry.get("extra", [])
            plan_n = len(plan)
            done_n = len([t for t in plan if t in done])
            total_plan += plan_n
            total_done += done_n
            projects_summary[proj] = {
                "plan": plan_n,
                "done": done_n,
                "extra": len(extra),
                "pct": round(done_n / plan_n * 100) if plan_n > 0 else 0
            }

        pct = round(total_done / total_plan * 100) if total_plan > 0 else 0

        result.append({
            "iso_week": iso_week,
            "week_start": week_start,
            "week_end": week_end,
            "plan_count": total_plan,
            "done_count": total_done,
            "pct": pct,
            "has_data": True,
            "projects": projects_summary
        })

    return result


# ── DOCX генерація ─────────────────────────────────────────────────────────────

def _generate_docx(iso_week: str) -> bytes:
    """Генерує DOCX-звіт у форматі план-факт."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx не встановлено. Запустіть: pip3 install python-docx")

    data = _load_json(_week_file(iso_week), None)
    projects = _load_json(PROJECTS_FILE, [])

    week_start, week_end = _week_bounds(iso_week)
    start_ua = _format_date_ua(week_start)
    end_ua = _format_date_ua(week_end)
    year = week_end[:4]
    period = f"{start_ua}-{end_ua}.{year}"

    doc = Document()

    # Поля сторінки
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Стиль за замовчуванням
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    def _add_centered(text, bold=False, size=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        return p

    def _set_table_borders(table):
        """Додає межі до таблиці."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)

    def _cell_text(cell, text, bold=False, size=11):
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'

    def _fill_done_cell(cell, done_tasks, comments):
        """Колонка «Виконана робота»: зелена ✓ ВИКОНАНО перед кожним пунктом."""
        cell.text = ''
        for idx, t in enumerate(done_tasks):
            p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            # Префікс ✓ ВИКОНАНО: — зелений жирний
            run_prefix = p.add_run("✓ ")
            run_prefix.bold = True
            run_prefix.font.size = Pt(11)
            run_prefix.font.name = 'Times New Roman'
            run_prefix.font.color.rgb = RGBColor(0x1E, 0x8A, 0x44)
            # Текст задачі — звичайний
            run_task = p.add_run(t)
            run_task.font.size = Pt(11)
            run_task.font.name = 'Times New Roman'
            # Коментар — курсив, сірий
            if t in comments:
                p_c = cell.add_paragraph()
                run_c = p_c.add_run(f"    → {comments[t]}")
                run_c.italic = True
                run_c.font.size = Pt(10)
                run_c.font.name = 'Times New Roman'
                run_c.font.color.rgb = RGBColor(0x55, 0x77, 0x99)

    # ── Заголовок ──────────────────────────────────────────────────────────────
    _add_centered("З В І Т – Жутов О.М", bold=True, size=14)
    _add_centered(f"Звітний період: {period}", bold=False, size=12)
    doc.add_paragraph()

    # ── Таблиця ПЛАН-ФАКТ ──────────────────────────────────────────────────────
    _add_centered("ПЛАН-ФАКТ", bold=True, size=12)

    plan_fact_table = doc.add_table(rows=1, cols=3)
    plan_fact_table.style = 'Table Grid'
    _set_table_borders(plan_fact_table)

    # Ширина колонок
    widths = [Inches(1.6), Inches(2.7), Inches(2.7)]
    for i, cell in enumerate(plan_fact_table.rows[0].cells):
        cell.width = widths[i]

    # Заголовки
    headers = ["Назва проекту", "План", "Виконана робота"]
    for i, cell in enumerate(plan_fact_table.rows[0].cells):
        _cell_text(cell, headers[i], bold=True)

    # Рядки проектів
    for proj in projects:
        entry = data.get("entries", {}).get(proj, {}) if data else {}
        cancelled = set(entry.get("cancelled", []))
        comments  = entry.get("comments", {})
        plan_tasks = [t for t in entry.get("plan", []) if t not in cancelled]
        done_tasks = [t for t in entry.get("done", []) if t not in cancelled]

        plan_text = "\n".join(f"- {t}" for t in plan_tasks) if plan_tasks else ""

        row = plan_fact_table.add_row()
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
        _cell_text(row.cells[0], proj)
        _cell_text(row.cells[1], plan_text)
        _fill_done_cell(row.cells[2], done_tasks, comments)

    doc.add_paragraph()

    # ── Таблиця ПОЗАПЛАН ───────────────────────────────────────────────────────
    _add_centered("ПОЗАПЛАН", bold=True, size=12)

    extra_table = doc.add_table(rows=1, cols=3)
    extra_table.style = 'Table Grid'
    _set_table_borders(extra_table)

    for i, cell in enumerate(extra_table.rows[0].cells):
        cell.width = widths[i]

    extra_headers = ["Назва проекту", "Виконана робота", "Планується"]
    for i, cell in enumerate(extra_table.rows[0].cells):
        _cell_text(cell, extra_headers[i], bold=True)

    has_extra = False
    if data:
        for proj in projects:
            entry = data.get("entries", {}).get(proj, {})
            extra = entry.get("extra", [])
            if extra:
                has_extra = True
                extra_text = "\n".join(f"- {t}" for t in extra)
                row = extra_table.add_row()
                for i, cell in enumerate(row.cells):
                    cell.width = widths[i]
                _cell_text(row.cells[0], proj)
                _cell_text(row.cells[1], extra_text)
                _cell_text(row.cells[2], "")

    if not has_extra:
        row = extra_table.add_row()
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    # ── Зберегти у байти ───────────────────────────────────────────────────────
    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HTTP обробник ──────────────────────────────────────────────────────────────

class Handler(BaseHTTPRequestHandler):

    def log_message(self, format, *args):
        """Логування запитів."""
        print(f"  {self.address_string()} → {format % args}")

    def _send_json(self, data, status=200):
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", len(body))
        self.end_headers()
        self.wfile.write(body)

    def _send_bytes(self, data, content_type, filename=None):
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", len(data))
        if filename:
            self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
        self.end_headers()
        self.wfile.write(data)

    def _send_error(self, msg, status=400):
        self._send_json({"error": msg}, status)

    def _read_body(self) -> dict:
        length = int(self.headers.get("Content-Length", 0))
        if length == 0:
            return {}
        raw = self.rfile.read(length)
        return json.loads(raw.decode("utf-8"))

    def do_GET(self):
        parsed = urlparse(self.path)
        path = parsed.path
        qs = parse_qs(parsed.query)

        # ── Головна сторінка ──────────────────────────────────────────────────
        if path in ("/", "/index.html", "/v1", "/v2"):
            file = HTML_FILE if path == "/v1" else HTML_FILE_V2
            with open(file, "rb") as f:
                body = f.read()
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", len(body))
            self.end_headers()
            self.wfile.write(body)
            return

        # ── API: тижень ───────────────────────────────────────────────────────
        m = re.match(r"^/api/week/(\d{4}-W\d{2})$", path)
        if m:
            iso_week = m.group(1)
            week_start, week_end = _week_bounds(iso_week)
            file_exists = _week_file(iso_week).exists()
            data = _load_json(_week_file(iso_week), {
                "iso_week": iso_week,
                "week_start": week_start,
                "week_end": week_end,
                "entries": {}
            })
            data["is_new"] = not file_exists
            self._send_json(data)
            return

        # ── API: історія ──────────────────────────────────────────────────────
        if path == "/api/history":
            n = int(qs.get("n", ["12"])[0])
            self._send_json(_get_history(n))
            return

        # ── API: проекти ──────────────────────────────────────────────────────
        if path == "/api/projects":
            projects = _load_json(PROJECTS_FILE, [])
            self._send_json(projects)
            return

        # ── API: поточний тиждень ─────────────────────────────────────────────
        if path == "/api/current-week":
            self._send_json({"iso_week": _current_iso_week()})
            return

        # ── API: експорт DOCX ─────────────────────────────────────────────────
        m = re.match(r"^/api/export/(\d{4}-W\d{2})$", path)
        if m:
            iso_week = m.group(1)
            week_start, week_end = _week_bounds(iso_week)
            start_ua = _format_date_ua(week_start).replace(".", "")
            end_ua = _format_date_ua(week_end).replace(".", "")
            filename = f"Zvit_{start_ua}-{end_ua}.docx"
            try:
                docx_bytes = _generate_docx(iso_week)
                self._send_bytes(
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    filename
                )
            except RuntimeError as e:
                self._send_error(str(e))
            return

        self._send_error("Not found", 404)

    def do_POST(self):
        parsed = urlparse(self.path)
        path = parsed.path

        # ── API: зберегти тиждень ─────────────────────────────────────────────
        m = re.match(r"^/api/week/(\d{4}-W\d{2})$", path)
        if m:
            iso_week = m.group(1)
            body = self._read_body()
            week_start, week_end = _week_bounds(iso_week)
            body["iso_week"] = iso_week
            body["week_start"] = week_start
            body["week_end"] = week_end
            _save_json(_week_file(iso_week), body)
            self._send_json({"ok": True})
            return

        # ── API: зберегти проекти ─────────────────────────────────────────────
        if path == "/api/projects":
            body = self._read_body()
            if isinstance(body, list):
                _save_json(PROJECTS_FILE, body)
                self._send_json({"ok": True})
            else:
                self._send_error("Expected array")
            return

        self._send_error("Not found", 404)

    def do_OPTIONS(self):
        self.send_response(204)
        self.end_headers()


# ── Запуск ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    port = 8080
    server = HTTPServer(("localhost", port), Handler)
    print(f"✅ Сервер запущено: http://localhost:{port}")
    print(f"   Дані: {DATA_DIR}")
    print("   Зупинка: Ctrl+C")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\n⛔ Сервер зупинено.")
